"""Document Preparation Module - Generate tender bidding documents"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from jinja2 import Template
from datetime import datetime
import os

class DocumentGenerator:
    """Generate professional tender documents"""
    
    def __init__(self):
        self.templates_dir = "/app/backend/document_templates"
        os.makedirs(self.templates_dir, exist_ok=True)
    
    def generate_boq(self, tender_data, items):
        """
        Generate Bill of Quantities (BOQ) Excel file
        
        Args:
            tender_data: Tender information
            items: List of BOQ items
        
        Returns:
            Path to generated Excel file
        """
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "BOQ"
        
        # Set column widths
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 15
        ws.column_dimensions['F'].width = 15
        
        # Header styling
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=12)
        
        # Title
        ws['A1'] = f"Bill of Quantities - {tender_data.get('tender_id', 'BOQ')}"
        ws['A1'].font = Font(bold=True, size=14)
        ws.merge_cells('A1:F1')
        
        # Tender details
        ws['A2'] = f"Tender: {tender_data.get('title', '')}"
        ws.merge_cells('A2:F2')
        ws['A3'] = f"Organization: {tender_data.get('organization', '')}"
        ws.merge_cells('A3:F3')
        ws['A4'] = f"Date: {datetime.now().strftime('%d-%m-%Y')}"
        ws.merge_cells('A4:F4')
        
        # Column headers
        headers = ['S.No', 'Description', 'Unit', 'Quantity', 'Rate (₹)', 'Amount (₹)']
        for col, header in enumerate(headers, start=1):
            cell = ws.cell(row=6, column=col)
            cell.value = header
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # BOQ items
        row = 7
        total_amount = 0
        
        for idx, item in enumerate(items, start=1):
            ws.cell(row=row, column=1, value=idx)
            ws.cell(row=row, column=2, value=item.get('description', ''))
            ws.cell(row=row, column=3, value=item.get('unit', 'Nos'))
            ws.cell(row=row, column=4, value=item.get('quantity', 1))
            ws.cell(row=row, column=5, value=item.get('rate', 0))
            
            amount = item.get('quantity', 1) * item.get('rate', 0)
            ws.cell(row=row, column=6, value=amount)
            ws.cell(row=row, column=6).number_format = '₹#,##0.00'
            
            total_amount += amount
            row += 1
        
        # Total row
        ws.cell(row=row, column=5, value="Total:")
        ws.cell(row=row, column=5).font = Font(bold=True)
        ws.cell(row=row, column=6, value=total_amount)
        ws.cell(row=row, column=6).font = Font(bold=True)
        ws.cell(row=row, column=6).number_format = '₹#,##0.00'
        
        # GST row
        row += 1
        gst_amount = total_amount * 0.18
        ws.cell(row=row, column=5, value="GST @ 18%:")
        ws.cell(row=row, column=6, value=gst_amount)
        ws.cell(row=row, column=6).number_format = '₹#,##0.00'
        
        # Grand total
        row += 1
        grand_total = total_amount + gst_amount
        ws.cell(row=row, column=5, value="Grand Total:")
        ws.cell(row=row, column=5).font = Font(bold=True, size=12)
        ws.cell(row=row, column=6, value=grand_total)
        ws.cell(row=row, column=6).font = Font(bold=True, size=12)
        ws.cell(row=row, column=6).number_format = '₹#,##0.00'
        
        # Save file
        filename = f"BOQ_{tender_data.get('tender_id', 'document')}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        filepath = os.path.join(self.templates_dir, filename)
        wb.save(filepath)
        
        return filepath
    
    def generate_cover_letter(self, company_data, tender_data):
        """
        Generate cover letter for tender submission
        
        Args:
            company_data: Company information
            tender_data: Tender information
        
        Returns:
            Path to generated Word document
        """
        doc = Document()
        
        # Company letterhead (placeholder)
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = company_data.get('name', 'Company Name')
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.runs[0].font.size = Pt(16)
        header_para.runs[0].font.bold = True
        
        # Date
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")
        doc.add_paragraph()
        
        # To address
        to_para = doc.add_paragraph()
        to_para.add_run(f"To,\n{tender_data.get('organization', 'Organization')}\n").bold = True
        
        # Subject
        doc.add_paragraph()
        subject = doc.add_paragraph()
        subject.add_run(f"Subject: Submission of Bid for {tender_data.get('tender_id', 'Tender ID')}\n").bold = True
        
        # Salutation
        doc.add_paragraph("Dear Sir/Madam,")
        doc.add_paragraph()
        
        # Body
        body_text = f"""With reference to your tender notice {tender_data.get('tender_id', '')}, we hereby submit our bid for "{tender_data.get('title', '')}".

We, {company_data.get('name', 'Company Name')}, have carefully studied the tender documents and confirm that we meet all the eligibility criteria and technical specifications mentioned therein.

We have attached the following documents:
1. Technical Bid
2. Financial Bid (BOQ)
3. Company Profile
4. Experience Certificates
5. EMD/Earnest Money Deposit
6. Required Certificates and Compliance Documents

We assure you of our best services and timely execution of the work/supply as per the terms and conditions of the tender.

We look forward to a positive response from your esteemed organization.

Thanking you,
"""
        doc.add_paragraph(body_text)
        
        # Signature
        doc.add_paragraph()
        signature = doc.add_paragraph("Yours faithfully,\n\n")
        doc.add_paragraph(f"{company_data.get('authorized_person', 'Authorized Signatory')}")
        doc.add_paragraph(f"{company_data.get('designation', 'Director')}")
        doc.add_paragraph(f"{company_data.get('name', 'Company Name')}")
        doc.add_paragraph(f"Contact: {company_data.get('phone', '')}")
        doc.add_paragraph(f"Email: {company_data.get('email', '')}")
        
        # Save file
        filename = f"CoverLetter_{tender_data.get('tender_id', 'document')}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(self.templates_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_company_profile(self, company_data):
        """
        Generate company profile document
        
        Args:
            company_data: Company information
        
        Returns:
            Path to generated Word document
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('COMPANY PROFILE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Company details
        sections = [
            ('Company Name', company_data.get('name', '')),
            ('Year of Establishment', company_data.get('established_year', '')),
            ('Registration Number', company_data.get('registration_number', '')),
            ('GST Number', company_data.get('gst_number', '')),
            ('PAN Number', company_data.get('pan_number', '')),
            ('Registered Address', company_data.get('address', '')),
            ('Contact Number', company_data.get('phone', '')),
            ('Email', company_data.get('email', '')),
            ('Website', company_data.get('website', '')),
        ]
        
        for label, value in sections:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)
        
        doc.add_paragraph()
        
        # About section
        doc.add_heading('About Us', 1)
        doc.add_paragraph(company_data.get('about', 'Company description here...'))
        
        # Services
        doc.add_heading('Our Services', 1)
        services = company_data.get('services', ['Service 1', 'Service 2', 'Service 3'])
        for service in services:
            doc.add_paragraph(service, style='List Bullet')
        
        # Key projects
        doc.add_heading('Key Projects', 1)
        projects = company_data.get('projects', [])
        for idx, project in enumerate(projects, start=1):
            doc.add_paragraph(f"{idx}. {project}")
        
        # Certifications
        doc.add_heading('Certifications', 1)
        certifications = company_data.get('certifications', ['ISO 9001:2015', 'ISO 27001'])
        for cert in certifications:
            doc.add_paragraph(cert, style='List Bullet')
        
        # Save file
        filename = f"CompanyProfile_{company_data.get('name', 'company').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(self.templates_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_technical_bid(self, tender_data, technical_details):
        """
        Generate technical bid document
        
        Args:
            tender_data: Tender information
            technical_details: Technical specifications and compliance
        
        Returns:
            Path to generated Word document
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('TECHNICAL BID', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Tender details
        doc.add_heading('Tender Details', 1)
        details = [
            ('Tender ID', tender_data.get('tender_id', '')),
            ('Tender Title', tender_data.get('title', '')),
            ('Organization', tender_data.get('organization', '')),
            ('Submission Date', datetime.now().strftime('%d-%m-%Y')),
        ]
        
        for label, value in details:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)
        
        doc.add_paragraph()
        
        # Eligibility compliance
        doc.add_heading('Eligibility Criteria Compliance', 1)
        doc.add_paragraph("We hereby confirm compliance with all eligibility criteria:")
        
        for criterion in tender_data.get('eligibility_criteria', []):
            doc.add_paragraph(f"✓ {criterion}", style='List Bullet')
        
        # Technical specifications
        doc.add_heading('Technical Specifications', 1)
        doc.add_paragraph("We confirm that our offered solution meets all technical requirements:")
        
        for spec_key, spec_value in technical_details.items():
            p = doc.add_paragraph()
            p.add_run(f"{spec_key}: ").bold = True
            p.add_run(str(spec_value))
        
        # Compliance matrix
        doc.add_heading('Compliance Matrix', 1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Medium Shading 1 Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Requirement'
        header_cells[1].text = 'Compliance'
        header_cells[2].text = 'Remarks'
        
        # Data rows
        for criterion in tender_data.get('eligibility_criteria', []):
            row_cells = table.add_row().cells
            row_cells[0].text = criterion
            row_cells[1].text = 'Yes'
            row_cells[2].text = 'Fully compliant'
        
        # Declaration
        doc.add_paragraph()
        doc.add_heading('Declaration', 1)
        declaration_text = """We declare that all information provided in this technical bid is true and accurate to the best of our knowledge. We understand that any false information may lead to disqualification."""
        doc.add_paragraph(declaration_text)
        
        # Save file
        filename = f"TechnicalBid_{tender_data.get('tender_id', 'document')}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(self.templates_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def calculate_emd(self, tender_value, emd_percentage=2.0):
        """
        Calculate EMD (Earnest Money Deposit)
        
        Args:
            tender_value: Estimated tender value
            emd_percentage: EMD percentage (default 2%)
        
        Returns:
            EMD amount
        """
        return (tender_value * emd_percentage) / 100
    
    def calculate_security_deposit(self, tender_value, sd_percentage=10.0):
        """
        Calculate Security Deposit
        
        Args:
            tender_value: Tender value
            sd_percentage: SD percentage (default 10%)
        
        Returns:
            Security deposit amount
        """
        return (tender_value * sd_percentage) / 100
